from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import time
from bs4 import BeautifulSoup
import re
import docx
from docx.shared import Inches
import requests
import shutil
import os

com = 'amazon'
exam = [
['ans-c00','AWS Certified Advanced Networking - Specialty',36,359],
['aws-certified-alexa-skill-builder-specialty','AWS Certified Alexa Skill Builder - Specialty',15,59],
['aws-certified-big-data-specialty','AWS Certified Big Data - Specialty Exam',22,85],
['aws-certified-cloud-practitioner','AWS Certified Cloud Practitioner (CLF-C01)',68,671],
['aws-certified-data-analytics-specialty','AWS Certified Data Analytics - Specialty (DAS-C01)',26,103],
['aws-certified-database-specialty','AWS Certified Database - Specialty (beta)',37,145],
['aws-certified-developer-associate','AWS Certified Developer Associate',37,367],
['aws-certified-machine-learning-specialty','AWS Certified Machine Learning - Specialty (MLS-C01)',33,131],
['aws-certified-security-specialty','AWS Certified Security - Specialty',50,249],
['aws-certified-solutions-architect-associate','AWS Certified Solutions Architect - Associate (SAA-C01)',35,346],
['aws-certified-solutions-architect-associate-saa-c02','AWS Certified Solutions Architect - Associate SAA-C02',46,457],
['aws-certified-solutions-architect-professional','AWS Certified Solutions Architect - Professional',75,750],
['aws-certified-sysops-administrator-associate','AWS Certified SysOps Administrator - Associate (SOA-C02)',14,54],
['aws-devops-engineer-professional','AWS DevOps Engineer - Professional (DOP-C01)',54,536],
['aws-sysops','AWS Certified SysOps Administrator',94,932]
]

chrome_options = Options()
chrome_options.add_argument("--headless")
driver = webdriver.Chrome('C:/Mine/code/chromedriver',options=chrome_options)
exam_total = open('C:/Mine/aws_exam/'+com+'_total.txt','a')
for i in range(0,len(exam)):
	num = 1
	q_num = 1
	question_file = docx.Document()
	discussion_file = docx.Document()
	ans_file = open('C:/Mine/aws_exam/'+exam[i][1]+'_ans.csv','a')
	error_file = open('C:/Mine/aws_exam/'+exam[i][1]+'_error.txt','a')
	delete_picture = []
	while num <= exam[i][2]:
		driver.get('http://www.google.com')
		time.sleep(3)
		exam_url = 'https://www.examtopics.com/exams/'+com+'/'+exam[i][0]+'/view/'+str(num)+'/\n'
		driver.find_element_by_class_name('gLFyf.gsfi').send_keys(exam_url)
		time.sleep(3)
		try:
			driver.find_element_by_class_name('action-menu').click()
			time.sleep(3)
			driver.find_element_by_class_name('fl').click()
			time.sleep(3)
		except:
			error_file.write(exam_url)
			num += 1
			continue
		soup = BeautifulSoup(driver.page_source,'html.parser')
		try:
			ans = soup.find_all('span',class_='correct-answer')
			for a in range(len(ans)):
				ans[a] = ans[a].text
			question_num = soup.find_all('div',class_='card-header text-white bg-primary')
			for j in range(len(question_num)):
				temp1 = question_num[j].text.split('T')
				question_num[j] = [temp1[0].strip(),'T'+temp1[-1].strip()]
			question = soup.find_all('p',class_='card-text')
			for a in range(len(question)):
				if question[a].find_all('img',class_='in-exam-image'):
					img = question[a].find_all('img',class_='in-exam-image')
					question[a] = [question[a].text.strip()]
					for d in img:
						question[a].append(d.get('src'))
				else:
					question[a] = [question[a].text.strip()]
			choices = soup.find_all('div',class_='question-choices-container')
			for a in range(len(choices)):
				c_temp = choices[a].find_all('li',class_='multi-choice-item')
				c_temp_num = 0
				for b in range(len(c_temp)):
					c_temp[b] = chr(ord('A')+c_temp_num)+'. '+c_temp[b].text.split(chr(ord('A')+c_temp_num)+'.')[-1].strip()
					c_temp_num += 1
				choices[a] = c_temp
			discussion = soup.find_all('a', class_='btn btn-secondary question-discussion-button d-print-none')
			driver.execute_script("window.open('');")
			time.sleep(3)
			driver.switch_to.window(driver.window_handles[1])
			time.sleep(3)
			run_num = 0
			run_num2 = 0
			for k in discussion:
				try:
					d_temp = k.get('href')
					print(exam[i][1])
					print('page ',num)
					print(question_num[run_num])
					question_file.add_paragraph(question_num[run_num][1]+' '+question_num[run_num][0]+'\n')
					if len(question[run_num2]) > 1:
						for f in range(1,len(question[run_num2])):
							r = requests.get('https://www.examtopics.com'+question[run_num2][f],stream = True)
							filename = ('https://www.examtopics.com'+question[run_num2][f]).split('/')[-1]
							r.raw.decode_content = True
							with open('C:/Mine/aws_exam/'+filename,'wb') as g:
								shutil.copyfileobj(r.raw, g)
							question_file.add_picture('C:/Mine/aws_exam/'+filename, width=Inches(5))
							delete_picture.append('C:/Mine/aws_exam/'+filename)
						question_file.add_paragraph(question[run_num2][0]+'\n')
					else:
						question_file.add_paragraph(question[run_num2][0]+'\n')
					for e in choices[run_num]:
						question_file.add_paragraph(str(e))
					question_file.add_paragraph('\n')
					ans_file.write(str(question_num[run_num][1]+','+question_num[run_num][0]+',,'+ans[run_num]+'\n'))
					driver.get('https://www.examtopics.com'+d_temp)
					time.sleep(5)
					d_soup = BeautifulSoup(driver.page_source,'html.parser')
					discussion_user = d_soup.find_all('h5',class_='mt-0 mb-0 align-middle d-inline-block comment-username')
					discussion_content = d_soup.find_all('div',class_='comment-content')
					discussion_file.add_paragraph(str(question_num[run_num][1]+' '+question_num[run_num][0]+'\n'))
					for c in range(len(discussion_content)):
						discussion_file.add_paragraph(discussion_user[c].text.strip())
						discussion_file.add_paragraph(discussion_content[c].text.strip())
					discussion_file.add_paragraph('\n')
					run_num += 1
					run_num2 += 2
					print('Finish\n')
					time.sleep(5)
				except:
					error_file.write(exam_url.strip()+'\t'+'\t'.join(question_num[run_num])+'\n')
					print('Error\n')
					run_num += 1
					run_num2 += 2
					continue
			driver.close()
			driver.switch_to.window(driver.window_handles[0])
			time.sleep(60)
		except:
			error_file.write(exam_url)
			num += 1
			continue
		num += 1
	question_file.save('C:/Mine/aws_exam/'+exam[i][1]+'_question.docx')
	discussion_file.save('C:/Mine/aws_exam/'+exam[i][1]+'_discussion.docx')
	for h in delete_picture:
		os.remove(h)
	ans_file.close()
	error_file.close()
	exam_total.write(exam[i][1]+'\t'+str(num-1)+'\n')
exam_total.close()